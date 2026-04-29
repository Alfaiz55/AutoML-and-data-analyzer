from db import init_db, save_result, get_history
init_db()

import streamlit as st
import pandas as pd
import tempfile
from pipe import run_pipeline

# =========================
# PAGE CONFIG (IMPORTANT)
# =========================
st.set_page_config(
    page_title="AutoML Analyzer",
    layout="centered"   # 🔥 mobile friendly (NOT wide)
)

# =========================
# 🔒 ADMIN PANEL
# =========================
st.sidebar.markdown("### 🔒 Admin Panel")

admin_mode = st.sidebar.checkbox("Enable Admin Mode")
admin_authenticated = False

if admin_mode:
    password = st.sidebar.text_input("Enter Admin Password", type="password")

    if password == "admin123":
        admin_authenticated = True
        st.sidebar.success("Access Granted")
    elif password:
        st.sidebar.error("Wrong Password")

# =========================
# 🎨 CLEAN MOBILE CSS
# =========================
st.markdown("""
<style>
/* Global */
body {
    background-color: #eef5ff;
}

/* Reduce padding for mobile */
.block-container {
    padding-top: 1rem;
    padding-bottom: 1rem;
    padding-left: 1rem;
    padding-right: 1rem;
}

/* Buttons full width */
button[kind="primary"] {
    width: 100% !important;
    border-radius: 10px !important;
}

/* Cards */
.card {
    background: white;
    padding: 14px;
    border-radius: 12px;
    box-shadow: 0px 3px 10px rgba(0,0,0,0.08);
    margin-bottom: 12px;
}

/* Headings */
h1, h2, h3 {
    color: #1f3c5b !important;
}

/* DataFrame scroll fix */
.stDataFrame {
    overflow-x: auto;
}
</style>
""", unsafe_allow_html=True)

# =========================
# HEADER
# =========================
st.title("🤖 AutoML Analyzer")
st.caption("Upload → Analyze → Get Results")

# =========================
# 📂 FILE SECTION
# =========================
with st.container():
    st.subheader("📂 Upload Dataset")

    uploaded_file = st.file_uploader(
        "Choose file",
        type=["csv", "xlsx", "xls", "docx"]
    )

df = None
target = None
valid_file = True

# =========================
# FILE PROCESSING
# =========================
if uploaded_file:

    file_ext = uploaded_file.name.split(".")[-1].lower()

    if file_ext not in ["csv", "xlsx", "xls", "docx"]:
        st.error("Unsupported file type.")
        valid_file = False

    elif file_ext == "docx":
        try:
            from docx import Document
            doc = Document(uploaded_file)

            if not doc.tables:
                st.error("No table found in Word file.")
                valid_file = False
            else:
                data = [[cell.text.strip() for cell in row.cells] for row in doc.tables[0].rows]
                df = pd.DataFrame(data[1:], columns=data[0])

        except ImportError:
            st.error("Install python-docx")
            valid_file = False

    elif file_ext == "csv":
        try:
            df = pd.read_csv(uploaded_file, encoding="utf-8")
        except:
            uploaded_file.seek(0)
            df = pd.read_csv(uploaded_file, encoding="latin1")

    else:
        df = pd.read_excel(uploaded_file)

# =========================
# DATA PREVIEW + OPTIONS
# =========================
if valid_file and df is not None:

    st.subheader("🔍 Data Preview")
    st.dataframe(df.head(), use_container_width=True)

    with st.expander("⚙️ Data Settings"):
        drop_cols = st.multiselect("Drop Columns", df.columns)

        if drop_cols:
            df = df.drop(columns=drop_cols)

        target = st.selectbox("Target Column", ["None"] + list(df.columns))
        if target == "None":
            target = None

# =========================
# 🚀 RUN BUTTON
# =========================
run = st.button("🚀 Start Analysis", use_container_width=True)

# =========================
# RESULT SECTION
# =========================
if uploaded_file and run and valid_file and df is not None:

    if target is None:
        st.info("Running Clustering Mode")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".csv") as tmp:
        df.to_csv(tmp.name, index=False)
        file_path = tmp.name

    result = run_pipeline(file_path, target)

    if "error" in result:
        st.error(result["error"])

    else:
        save_result(uploaded_file.name, result)

        # ===== RESULT CARD =====
        st.markdown("### 📊 Results")

        st.success(f"""
        **Problem Type:** {result['problem_type']}  
        **Best Model:** {result['best_model']}  
        **Score:** {round(result['score'], 4)}
        """)

        # ===== MODEL TABLE =====
        st.subheader("📈 Model Comparison")

        results = result.get("all_results", {})
        results = {k: v for k, v in results.items() if "BEST" not in k}

        if results:
            df_models = pd.DataFrame(results.items(), columns=["Model", "Score"])
            st.dataframe(df_models.sort_values(by="Score", ascending=False), use_container_width=True)

        # ===== INSIGHTS =====
        st.subheader("🧠 Insights")
        st.write(result.get("insights", ""))

        # ===== VISUALS =====
        st.subheader("📊 Visualizations")

        if result.get("plots"):
            for fig in result["plots"]:
                st.pyplot(fig, use_container_width=True)
        else:
            st.info("No plots generated")

# =========================
# ADMIN HISTORY
# =========================
if admin_authenticated:

    st.subheader("🧠 Admin History")

    history = get_history()

    if history:
        df_history = pd.DataFrame(history)

        if df_history.shape[1] == 6:
            df_history.columns = ["ID", "Dataset", "Problem Type", "Best Model", "Score", "Timestamp"]
            df_history = df_history.drop(columns=["ID"])
        else:
            df_history.columns = ["Dataset", "Problem Type", "Best Model", "Score", "Timestamp"]

        st.dataframe(df_history.sort_values(by="Timestamp", ascending=False), use_container_width=True)

    else:
        st.info("No history found.")
