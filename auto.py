# from db import init_db, save_result, get_history
# init_db()

# import streamlit as st
# import pandas as pd
# import tempfile
# from pipe import run_pipeline

# st.set_page_config(page_title="AutoML Analyzer", layout="wide")

# # =========================
# # 🔒 ADMIN PANEL (SIDEBAR)
# # =========================
# st.sidebar.markdown("### 🔒 Admin Panel")

# admin_mode = st.sidebar.checkbox("Enable Admin Mode")

# admin_authenticated = False

# if admin_mode:
#     password = st.sidebar.text_input("Enter Admin Password", type="password")

#     if password == "admin123":
#         admin_authenticated = True
#         st.sidebar.success("Access Granted")
#     elif password:
#         st.sidebar.error("Wrong Password")

# # =========================
# # 🎨 IMPROVED LIGHT BLUE UI
# # =========================
# st.markdown("""
# <style>
# .stApp {
#     background: linear-gradient(135deg, #e6f2ff, #f4f8fb);
# }

# .block-container {
#     padding-top: 2rem;
# }

# .card {
#     background: white;
#     padding: 20px;
#     border-radius: 16px;
#     box-shadow: 0px 6px 18px rgba(0,0,0,0.08);
#     margin-bottom: 20px;
# }

# .title {
#     font-size: 32px;
#     font-weight: bold;
#     color: #1f3c5b;
# }

# .subtitle {
#     color: #5d6d7e;
#     margin-bottom: 20px;
# }

# .section-title {
#     font-size: 18px;
#     font-weight: 600;
#     color: #2c3e50;
#     margin-bottom: 10px;
# }
# /* ===== FIX TEXT COLORS ===== */

# html, body, [class*="css"] {
#     color: #1f2d3d !important;
# }

# /* Headers */
# h1, h2, h3, h4, h5, h6 {
#     color: #1f3c5b !important;
# }

# /* Streamlit text */
# .stMarkdown, .stText {
#     color: #1f2d3d !important;
# }

# /* Labels */
# label {
#     color: #2c3e50 !important;
# }

# /* Dataframe text */
# .stDataFrame {
#     color: black !important;
# }     
# div[data-testid="stAlert"] {
#     background-color: #fff3cd !important;
#     border-left: 6px solid #ffcc00 !important;
#     color: black !important;
#     border-radius: 8px;
# }                   
# </style>
# """, unsafe_allow_html=True)

# # =========================
# # HEADER
# # =========================
# st.markdown('<div class="title">🤖 AI AutoML Analyzer</div>', unsafe_allow_html=True)
# st.markdown('<div class="subtitle">Smart Machine Learning Analysis</div>', unsafe_allow_html=True)

# # =========================
# # UPLOAD CARD
# # =========================
# st.markdown('<div class="card">', unsafe_allow_html=True)

# uploaded_file = st.file_uploader(
#     "📂 Upload Dataset",
#     type=["csv", "xlsx", "xls", "docx"]   # ✅ added formats
# )

# target = None

# if uploaded_file:

#     # =========================
#     # 📄 DOCX HANDLING
#     # =========================
#     if uploaded_file.name.endswith(".docx"):

#         from docx import Document

#         def docx_to_df(file):
#             doc = Document(file)

#             if not doc.tables:
#                 return None

#             data = []
#             for row in doc.tables[0].rows:
#                 data.append([cell.text.strip() for cell in row.cells])

#             return pd.DataFrame(data[1:], columns=data[0])

#         df = docx_to_df(uploaded_file)

#         if df is None:
#             st.warning("⚠️ No table found in Word file. Please upload a valid dataset.")
#             st.stop()   # 🔥 stop execution safely

#     # =========================
#     # 📊 CSV HANDLING
#     # =========================
#     elif uploaded_file.name.endswith(".csv"):

#         try:
#             df = pd.read_csv(uploaded_file, encoding="utf-8")
#         except UnicodeDecodeError:
#             uploaded_file.seek(0)
#             try:
#                 df = pd.read_csv(uploaded_file, encoding="latin1")
#             except:
#                 uploaded_file.seek(0)
#                 df = pd.read_csv(uploaded_file, encoding="ISO-8859-1")

#     # =========================
#     # 📈 EXCEL HANDLING
#     # =========================
#     elif uploaded_file.name.endswith((".xlsx", ".xls")):

#         df = pd.read_excel(uploaded_file)

#     # =========================
#     # 👀 OPTIONAL PREVIEW
#     # =========================
#     st.write("Preview of dataset:")
#     st.dataframe(df.head())

#     # =========================
#     # CLEAN DATA
#     # =========================
#     st.markdown('<div class="section-title">🧹 Clean Dataset</div>', unsafe_allow_html=True)

#     drop_cols = st.multiselect("Select columns to drop", df.columns)

#     if drop_cols:
#         df = df.drop(columns=drop_cols)

#     # =========================
#     # TARGET SELECT
#     # =========================
#     target = st.selectbox("🎯 Select Target Column", ["None"] + list(df.columns))

#     if target == "None":
#         target = None

# run = st.button("🚀 Start Analysis")

# st.markdown('</div>', unsafe_allow_html=True)

# # =========================
# # RUN PIPELINE
# # =========================
# if uploaded_file and run:

#     if target is None:
#         st.info("⚙️ No target selected → Running Clustering Mode")

#     with st.spinner("Running AutoML Pipeline..."):

#         with tempfile.NamedTemporaryFile(delete=False, suffix=".csv") as tmp:
#             df.to_csv(tmp.name, index=False)
#             file_path = tmp.name

#         result = run_pipeline(file_path, target)

#     if "error" in result:
#         st.error(result["error"])

#     else:
#         save_result(uploaded_file.name, result)

#         # =========================
#         # RESULT CARDS
#         # =========================
#         col1, col2, col3 = st.columns(3)

#         with col1:
#             st.markdown('<div class="card">', unsafe_allow_html=True)
#             st.markdown("### 🧠 Problem Type")
#             st.write(result["problem_type"])
#             st.markdown('</div>', unsafe_allow_html=True)

#         with col2:
#             st.markdown('<div class="card">', unsafe_allow_html=True)
#             st.markdown("### 🤖 Best Model")
#             st.write(result["best_model"])
#             st.write(f"Score: {round(result['score'],4)}")
#             st.markdown('</div>', unsafe_allow_html=True)

#         with col3:
#             st.markdown('<div class="card">', unsafe_allow_html=True)
#             st.markdown("### 📊 Dataset Info")
#             st.write(f"Rows: {df.shape[0]}")
#             st.write(f"Columns: {df.shape[1]}")
#             st.write(f"Missing: {df.isnull().sum().sum()}")
#             st.markdown('</div>', unsafe_allow_html=True)

#         # =========================
#         # 📊 MODEL COMPARISON
#         # =========================
#         st.markdown('<div class="card">', unsafe_allow_html=True)
#         st.markdown("### 📊 Model Comparison")

#         all_results = result.get("all_results", {})

#         filtered_results = {
#             k: v for k, v in all_results.items()
#             if k not in ["BEST_MODEL", "BEST_SCORE"]
#         }

#         if filtered_results:
#             df_models = pd.DataFrame(
#                 list(filtered_results.items()),
#                 columns=["Model", "Score"]
#             )

#             df_models = df_models.sort_values(by="Score", ascending=False)

#             st.write(f"Total Models Tested: {len(df_models)}")
#             st.dataframe(df_models, use_container_width=True)

#         else:
#             st.info("No model comparison available")

#         st.markdown('</div>', unsafe_allow_html=True)

#         # =========================
#         # ANALYSIS SUMMARY
#         # =========================
#         st.markdown('<div class="card">', unsafe_allow_html=True)
#         st.markdown("### 📄 Analysis Overview")

#         if result.get("insights"):
#             st.write(result["insights"])
#         else:
#             st.write(f"- Problem Type: {result['problem_type']}")
#             st.write(f"- Best Model: {result['best_model']}")
#             st.write(f"- Score: {round(result['score'],4)}")

#         st.markdown('</div>', unsafe_allow_html=True)

     

#         # =========================
#         # VISUALIZATION
#         # =========================
#         st.markdown('<div class="card">', unsafe_allow_html=True)
#         st.markdown("### 📈 Data Visualizations")

#         if result.get("plots"):
#             cols = st.columns(2)

#             for i, fig in enumerate(result["plots"]):
#                 with cols[i % 2]:
#                     st.pyplot(fig)
#         else:
#             st.warning("No plots generated")

#         st.markdown('</div>', unsafe_allow_html=True)

# # =========================
# # 🔒 ADMIN HISTORY VIEW
# # =========================
# if admin_authenticated:

#     st.markdown("## 🧠 Admin History")

#     history = get_history()

#     if history:
#         df_history = pd.DataFrame(history, columns=[
#             "Dataset",
#             "Problem Type",
#             "Best Model",
#             "Score",
#             "Timestamp"
#         ])

#         df_history = df_history.sort_values(by="Timestamp", ascending=False)

#         st.dataframe(df_history, use_container_width=True)

#     else:
#         st.info("No history found.")


from db import init_db, save_result, get_history
init_db()

import streamlit as st
import pandas as pd
import tempfile
from pipe import run_pipeline

st.set_page_config(page_title="AutoML Analyzer", layout="wide")

# =========================
# 🔒 ADMIN PANEL
# =========================
st.sidebar.markdown("### 🔒 Admin Panel")

admin_mode = st.sidebar.checkbox("Enable Admin Mode")
admin_authenticated = False

if admin_mode:
    password = st.sidebar.text_input("Enter Admin Password", type="password")

    if password == "admin123":
        admin_authenticated = True
        st.sidebar.success("Access Granted")
    elif password:
        st.sidebar.error("Wrong Password")

# =========================
# 🎨 CLEAN UI STYLE (FIXED COLORS)
# =========================
st.markdown("""
<style>
html, body, p, div, span, label {
    color: #1a1a1a !important;
}

/* ===== SIDEBAR ===== */
section[data-testid="stSidebar"] {
    background-color: #1e1e2f !important;
}

section[data-testid="stSidebar"] * {
    color: #ffffff !important;
}

/* ===== MAIN BACKGROUND ===== */
.stApp {
    background-color: #eef5ff;
}

/* ===== CARDS ===== */
.card {
    background: white;
    padding: 18px;
    border-radius: 12px;
    box-shadow: 0px 4px 12px rgba(0,0,0,0.08);
}

/* ===== BUTTONS ===== */
button {
    color: white !important;
    background-color: #007bff !important;
    border-radius: 8px !important;
}

/* ===== TABLE (VERY IMPORTANT) ===== */
.stDataFrame table {
    background-color: white !important;
    color: black !important;
}

.stDataFrame th {
    background-color: #007bff !important;
    color: white !important;
    font-weight: bold;
}

.stDataFrame td {
    color: black !important;
}

/* ===== ALERTS ===== */
div[data-testid="stAlert"] {
    color: black !important;
    font-weight: 500;
    border-radius: 8px;
}

div[data-testid="stAlert"][kind="warning"] {
    background-color: #fff3cd !important;
    border-left: 6px solid #ffcc00 !important;
}

div[data-testid="stAlert"][kind="error"] {
    background-color: #f8d7da !important;
    border-left: 6px solid #dc3545 !important;
}

div[data-testid="stAlert"][kind="info"] {
    background-color: #d1ecf1 !important;
    border-left: 6px solid #17a2b8 !important;
}

/* ===== HEADINGS ===== */
h1, h2, h3 {
    color: #1f3c5b !important;
}

</style>
""", unsafe_allow_html=True)

# =========================
# HEADER
# =========================
st.title("🤖 AI AutoML Analyzer")
st.caption("Smart Machine Learning Analysis")

# =========================
# FILE UPLOAD
# =========================
st.markdown('<div class="card">', unsafe_allow_html=True)

uploaded_file = st.file_uploader(
    "📂 Upload Dataset",
    type=["csv", "xlsx", "xls", "docx"]
)

df = None
target = None
valid_file = True

if uploaded_file:

    file_ext = uploaded_file.name.split(".")[-1].lower()

    # ===== INVALID FILE =====
    if file_ext not in ["csv", "xlsx", "xls", "docx"]:
        st.error("Unsupported file type.")
        valid_file = False

    # ===== DOCX =====
    elif file_ext == "docx":
        try:
            from docx import Document
            doc = Document(uploaded_file)

            if not doc.tables:
                st.error("Word file does not contain a dataset table.")
                valid_file = False
            else:
                data = [[cell.text.strip() for cell in row.cells] for row in doc.tables[0].rows]
                df = pd.DataFrame(data[1:], columns=data[0])

        except ImportError:
            st.error("python-docx not installed. Run: pip install python-docx")
            valid_file = False

    # ===== CSV =====
    elif file_ext == "csv":
        try:
            df = pd.read_csv(uploaded_file, encoding="utf-8")
        except:
            uploaded_file.seek(0)
            df = pd.read_csv(uploaded_file, encoding="latin1")

    # ===== EXCEL =====
    else:
        df = pd.read_excel(uploaded_file)

    # ===== SHOW DATA ONLY IF VALID =====
    if valid_file and df is not None:

        st.subheader("Preview")
        st.dataframe(df.head())

        drop_cols = st.multiselect("Drop Columns", df.columns)
        if drop_cols:
            df = df.drop(columns=drop_cols)

        target = st.selectbox("Target Column", ["None"] + list(df.columns))
        if target == "None":
            target = None

run = st.button("🚀 Start Analysis")
st.markdown('</div>', unsafe_allow_html=True)

# =========================
# RUN MODEL
# =========================
if uploaded_file and run and valid_file and df is not None:

    if target is None:
        st.info("Running Clustering Mode")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".csv") as tmp:
        df.to_csv(tmp.name, index=False)
        file_path = tmp.name

    result = run_pipeline(file_path, target)

    if "error" in result:
        st.error(result["error"])
    else:
        save_result(uploaded_file.name, result)

        col1, col2, col3 = st.columns(3)
        col1.metric("Problem Type", result["problem_type"])
        col2.metric("Best Model", result["best_model"])
        col3.metric("Score", round(result["score"], 4))

        # MODEL COMPARISON
        st.subheader("📊 Model Comparison")

        results = result.get("all_results", {})
        results = {k: v for k, v in results.items() if "BEST" not in k}

        if results:
            df_models = pd.DataFrame(results.items(), columns=["Model", "Score"])
            st.dataframe(df_models.sort_values(by="Score", ascending=False))

        # INSIGHTS
        st.subheader("📄 Insights")
        st.write(result.get("insights", ""))

        # VISUALS
        st.subheader("📈 Visualizations")

        if result.get("plots"):
            for fig in result["plots"]:
                st.pyplot(fig)
        else:
            st.info("No plots generated")

# =========================
# ADMIN HISTORY (FIXED)
# =========================
if admin_authenticated:

    st.subheader("🧠 Admin History")

    history = get_history()

    if history:

        df_history = pd.DataFrame(history)

        # HANDLE BOTH CASES (5 or 6 columns)
        if df_history.shape[1] == 6:
            df_history.columns = ["ID", "Dataset", "Problem Type", "Best Model", "Score", "Timestamp"]
            df_history = df_history.drop(columns=["ID"])
        else:
            df_history.columns = ["Dataset", "Problem Type", "Best Model", "Score", "Timestamp"]

        st.dataframe(df_history.sort_values(by="Timestamp", ascending=False))

    else:
        st.info("No history found.")